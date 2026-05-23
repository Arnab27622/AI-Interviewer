import io
import re
import fitz  # PyMuPDF
from docx import Document
from app.services.extraction.ocr_service import OCRService
import logging

logger = logging.getLogger("FileProcessingService")

class FileProcessingService:
    @staticmethod
    def clean_text(text: str) -> str:
        """Removes excessive newlines and spaces from extracted text."""
        if not text:
            return ""
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.strip()
        return text

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        """Decodes standard text files with Unicode fallbacks."""
        logger.info("Parsing standard TXT file bytes")
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            text = file_bytes.decode('latin-1')
        return FileProcessingService.clean_text(text)

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """Extracts text from paragraphs and tables in DOCX documents."""
        logger.info("Parsing DOCX file paragraphs and tables")
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_data))
                
        return FileProcessingService.clean_text("\n".join(full_text))

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> tuple[str, str]:
        """
        Extracts text from PDF bytes via PyMuPDF.
        Triggers OCR fallback if text density is too low (< 100 chars per page).
        """
        logger.info("Parsing PDF file bytes via PyMuPDF")
        method = "pymupdf"
        text = ""
        
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text") + "\n"
            
            # Extract embedded hyperlinks so the LLM can parse them
            for link in page.get_links():
                if "uri" in link and link["uri"]:
                    text += f"{link['uri']}\n"
            
        cleaned_text = FileProcessingService.clean_text(text)
        
        if pdf_document.page_count > 0 and len(cleaned_text) < (pdf_document.page_count * 100):
            logger.info("PDF text density too low. Falling back to OCR processing.")
            method = "ocr"
            cleaned_text = OCRService.extract_text_from_pdf_bytes(file_bytes)
            cleaned_text = FileProcessingService.clean_text(cleaned_text)
            
        pdf_document.close()
        return cleaned_text, method

    @classmethod
    def process_file(cls, file_bytes: bytes, filename: str) -> dict:
        """Main router for parsing uploads based on file extension."""
        ext = filename.lower().split('.')[-1]
        logger.info(f"Processing uploaded file: '{filename}' with extension '{ext}'")
        
        if ext == 'txt':
            return {"raw_text": cls.parse_txt(file_bytes), "method_used": "txt"}
        elif ext == 'docx':
            return {"raw_text": cls.parse_docx(file_bytes), "method_used": "docx"}
        elif ext == 'pdf':
            text, method = cls.parse_pdf(file_bytes)
            return {"raw_text": text, "method_used": method}
        else:
            logger.error(f"Unsupported file format attempted: {ext}")
            raise ValueError(f"Unsupported file format: {ext}")
